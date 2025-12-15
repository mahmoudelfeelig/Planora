from __future__ import annotations
from typing import Dict, Any, Iterable, List, Optional
import csv
import re
from pathlib import Path
from datetime import date, time, datetime, timedelta, timezone

# Optional DOCX
try:
    from docx import Document
    from docx.shared import Pt
except Exception:  # pragma: no cover
    Document = None  # type: ignore

from domain import Instance

# -------- time labels --------

_DEFAULT_START = time(8, 30)
_DEFAULT_SLOT_MINUTES = 90
_DEFAULT_BREAK_MINUTES = 0

def _slot_labels(inst: Instance) -> List[str]:
    slots = inst.slots_per_day
    labels = getattr(inst, "time_labels", None)
    if isinstance(labels, list) and len(labels) == slots:
        return labels

    slot_minutes = getattr(inst, "slot_minutes", _DEFAULT_SLOT_MINUTES)
    day_start = getattr(inst, "day_start_time", None)
    if isinstance(day_start, str):
        try:
            hh, mm = map(int, day_start.split(":"))
            day_start = time(hh, mm)
        except Exception:
            day_start = _DEFAULT_START
    if not isinstance(day_start, time):
        day_start = _DEFAULT_START

    labels: List[str] = []
    cur = datetime.combine(date.today(), day_start)
    for _ in range(slots):
        start = cur
        end = start + timedelta(minutes=slot_minutes)
        labels.append(f"{start.strftime('%H:%M')} - {end.strftime('%H:%M')}")
        cur = end + timedelta(minutes=getattr(inst, "slot_break_minutes", _DEFAULT_BREAK_MINUTES))
    return labels

# -------- table helpers --------

def _iter_cells_for_view(
    inst: Instance,
    schedule: Dict[int, Dict[str, Any]],
    *,
    week: int,
    day: str,
    slot: int,
    view: str,
    entity_id: int,
) -> Iterable[int]:
    for a_id, info in schedule.items():
        if info["week"] != week or info["day"] != day:
            continue
        s0 = int(info["slot"])
        dur = int(info["duration"])
        if not (s0 <= slot < s0 + dur):
            continue
        if view == "group" and entity_id in info.get("group_ids", []):
            yield a_id
        elif view == "staff" and entity_id == info.get("staff_id"):
            yield a_id
        elif view == "room" and entity_id == info.get("room_id"):
            yield a_id

def _cell_text_for_activity(inst: Instance, schedule: Dict[int, Dict[str, Any]], a_id: int) -> str:
    info = schedule[a_id]
    course = inst.courses.get(info["course_id"])
    room = inst.rooms.get(info.get("room_id")) if info.get("room_id") is not None else None
    staff = inst.staff.get(info.get("staff_id")) if info.get("staff_id") is not None else None

    parts: List[str] = []
    if course:
        parts.append(course.code)
        parts.append(course.name)
    parts.append(str(info.get("kind", "")))
    if room:
        parts.append(f"Room: {room.name}")
    if staff:
        parts.append(f"Staff: {staff.name}")
    return "\n".join(parts)

def _ensure_docx():
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install 'python-docx' to enable DOCX exports.")

def _add_title(doc, text: str):
    h = doc.add_paragraph()
    r = h.add_run(text)
    r.font.size = Pt(16)
    r.bold = True

def _add_subtitle(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)

# -------- DOCX exports (unchanged behavior) --------

def export_group_schedules_to_docx(inst: Instance, schedule: Dict[int, Dict[str, Any]], filename: str) -> None:
    _ensure_docx()
    doc = Document()
    slot_labels = _slot_labels(inst)
    days = inst.days
    slots = inst.slots_per_day

    for g_id, g in inst.groups.items():
        _add_title(doc, f"Group: {g.name}")
        for w in inst.weeks:
            _add_subtitle(doc, f"Week {w}")
            table = doc.add_table(rows=len(days) + 1, cols=slots + 1)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Day / Time"
            for c in range(slots):
                table.cell(0, c + 1).text = slot_labels[c]
            for r, day in enumerate(days, start=1):
                table.cell(r, 0).text = day
                for c in range(slots):
                    cell = table.cell(r, c + 1)
                    entries = [
                        _cell_text_for_activity(inst, schedule, a)
                        for a in _iter_cells_for_view(inst, schedule, week=w, day=day, slot=c, view="group", entity_id=g_id)
                    ]
                    cell.text = "\n\n".join(entries) if entries else ""
        doc.add_page_break()

    doc.save(filename)

def export_staff_schedules_to_docx(inst: Instance, schedule: Dict[int, Dict[str, Any]], filename: str) -> None:
    _ensure_docx()
    doc = Document()
    slot_labels = _slot_labels(inst)
    days = inst.days
    slots = inst.slots_per_day

    for s_id, s in inst.staff.items():
        _add_title(doc, f"Staff: {s.name}")
        for w in inst.weeks:
            _add_subtitle(doc, f"Week {w}")
            table = doc.add_table(rows=len(days) + 1, cols=slots + 1)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Day / Time"
            for c in range(slots):
                table.cell(0, c + 1).text = slot_labels[c]
            for r, day in enumerate(days, start=1):
                table.cell(r, 0).text = day
                for c in range(slots):
                    cell = table.cell(r, c + 1)
                    entries = [
                        _cell_text_for_activity(inst, schedule, a)
                        for a in _iter_cells_for_view(inst, schedule, week=w, day=day, slot=c, view="staff", entity_id=s_id)
                    ]
                    cell.text = "\n\n".join(entries) if entries else ""
        doc.add_page_break()

    doc.save(filename)

def export_room_schedules_to_docx(inst: Instance, schedule: Dict[int, Dict[str, Any]], filename: str) -> None:
    _ensure_docx()
    doc = Document()
    slot_labels = _slot_labels(inst)
    days = inst.days
    slots = inst.slots_per_day

    for r_id, r in inst.rooms.items():
        _add_title(doc, f"Room: {r.name}")
        for w in inst.weeks:
            _add_subtitle(doc, f"Week {w}")
            table = doc.add_table(rows=len(days) + 1, cols=slots + 1)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Day / Time"
            for c in range(slots):
                table.cell(0, c + 1).text = slot_labels[c]
            for rr, day in enumerate(days, start=1):
                table.cell(rr, 0).text = day
                for cc in range(slots):
                    cell = table.cell(rr, cc + 1)
                    entries = [
                        _cell_text_for_activity(inst, schedule, a)
                        for a in _iter_cells_for_view(inst, schedule, week=w, day=day, slot=cc, view="room", entity_id=r_id)
                    ]
                    cell.text = "\n\n".join(entries) if entries else ""
        doc.add_page_break()

    doc.save(filename)

# -------- CSV export --------

def export_schedule_to_csv(inst: Instance, schedule: Dict[int, Dict[str, Any]], filename: str) -> None:
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["activity_id", "week", "day", "slot", "duration", "room_id", "staff_id", "course_id", "group_ids", "kind"])
        for a_id, info in schedule.items():
            gid_list = info.get("group_ids", [])
            gid_str = ";".join(str(x) for x in gid_list) if gid_list else ""
            w.writerow([
                a_id,
                info.get("week"),
                info.get("day"),
                info.get("slot"),
                info.get("duration"),
                info.get("room_id"),
                info.get("staff_id"),
                info.get("course_id"),
                gid_str,
                info.get("kind"),
            ])

# -------- Minimal PDF export (no external dependencies) --------

def _pdf_escape(s: str) -> str:
    return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_simple_pdf(pages: List[List[str]], out_path: str | Path) -> None:
    """
    Write a basic text-only PDF with one or more pages.
    This avoids external PDF dependencies and is sufficient for quick reporting.
    """
    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    objects: List[bytes] = []

    def add_obj(data: str) -> int:
        objects.append(data.encode("latin-1", errors="replace"))
        return len(objects)

    font_obj = add_obj("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_obj_ids: List[int] = []
    for page_lines in pages:
        x0 = 50
        y0 = 800
        leading = 12
        lines = ["BT", "/F1 10 Tf", f"{leading} TL", f"{x0} {y0} Td"]
        for line in page_lines:
            lines.append(f"({_pdf_escape(line)}) Tj")
            lines.append("T*")
        lines.append("ET")
        stream = "\n".join(lines) + "\n"

        content = f"<< /Length {len(stream.encode('latin-1', errors='replace'))} >>\nstream\n{stream}endstream"
        content_id = add_obj(content)

        page_dict = f"<< /Type /Page /Parent 0 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_obj} 0 R >> >> /Contents {content_id} 0 R >>"
        page_id = add_obj(page_dict)
        page_obj_ids.append(page_id)

    kids = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
    pages_id = add_obj(f"<< /Type /Pages /Count {len(page_obj_ids)} /Kids [ {kids} ] >>")

    for page_id in page_obj_ids:
        raw = objects[page_id - 1].decode("latin-1", errors="replace")
        objects[page_id - 1] = raw.replace("/Parent 0 0 R", f"/Parent {pages_id} 0 R").encode("latin-1", errors="replace")

    catalog_id = add_obj(f"<< /Type /Catalog /Pages {pages_id} 0 R >>")

    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    offsets = [0]
    body = b""
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(header) + len(body))
        body += f"{i} 0 obj\n".encode("latin-1") + obj + b"\nendobj\n"

    xref_start = len(header) + len(body)
    xref = [f"xref\n0 {len(objects)+1}\n".encode("latin-1")]
    xref.append(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        xref.append(f"{off:010d} 00000 n \n".encode("latin-1"))
    xref_bytes = b"".join(xref)

    trailer = (
        f"trailer\n<< /Size {len(objects)+1} /Root {catalog_id} 0 R >>\n"
        f"startxref\n{xref_start}\n%%EOF\n"
    ).encode("latin-1")

    path.write_bytes(header + body + xref_bytes + trailer)


def export_groups_pdf(inst: Instance, schedule: Dict[int, Dict[str, Any]], out_path: str | Path) -> None:
    """
    Text-only PDF export: one page per group with a sorted activity list.
    """
    pages: List[List[str]] = []
    day_order = {d: i for i, d in enumerate(inst.days)}

    for g_id, g in sorted(inst.groups.items()):
        lines: List[str] = [f"Group: {g.name} (id {g_id})"]
        items = []
        for a_id, info in schedule.items():
            if g_id not in info.get("group_ids", []):
                continue
            course = inst.courses.get(info["course_id"])
            room = inst.rooms.get(info.get("room_id")) if info.get("room_id") is not None else None
            staff = inst.staff.get(info.get("staff_id")) if info.get("staff_id") is not None else None
            items.append((
                int(info["week"]),
                day_order.get(str(info["day"]), 999),
                int(info["slot"]),
                a_id,
                course.code if course else str(info.get("course_id")),
                str(info.get("kind", "")),
                room.name if room else str(info.get("room_id")),
                staff.name if staff else str(info.get("staff_id")),
                int(info.get("duration", 1)),
                str(info.get("day")),
            ))
        items.sort()

        for (w, _, slot, a_id, code, kind, room_name, staff_name, dur, day) in items:
            lines.append(f"W{w:02d} {day} s{slot+1} dur{dur}  {code} {kind}  room={room_name}  staff={staff_name}  (A{a_id})")
        if len(lines) == 1:
            lines.append("(no scheduled activities)")
        pages.append(lines)

    _write_simple_pdf(pages, out_path)


# -------- richer reporting --------

def export_summary_reports(inst: Instance, schedule: Dict[int, Dict[str, Any]], out_dir: str | Path) -> None:
    """
    Write small CSV reports that are easy to inspect:
      - staff_load.csv: per staff/week total slots
      - group_load.csv: per group/week total slots
      - room_util.csv: per room/week total slots
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    staff_load: Dict[tuple[int, int], int] = {}
    group_load: Dict[tuple[int, int], int] = {}
    room_load: Dict[tuple[int, int], int] = {}

    for info in schedule.values():
        w = int(info["week"])
        dur = int(info["duration"])

        sid = info.get("staff_id")
        if sid is not None:
            key = (int(sid), w)
            staff_load[key] = staff_load.get(key, 0) + dur

        rid = info.get("room_id")
        if rid is not None:
            key = (int(rid), w)
            room_load[key] = room_load.get(key, 0) + dur

        for g in info.get("group_ids", []) or []:
            key = (int(g), w)
            group_load[key] = group_load.get(key, 0) + dur

    def _write(path: Path, header: List[str], rows: List[List[object]]) -> None:
        with path.open("w", newline="", encoding="utf-8") as f:
            cw = csv.writer(f)
            cw.writerow(header)
            cw.writerows(rows)

    _write(
        out_dir / "staff_load.csv",
        ["staff_id", "staff_name", "week", "slots"],
        [
            [sid, inst.staff.get(sid).name if sid in inst.staff else "", w, slots]
            for (sid, w), slots in sorted(staff_load.items())
        ],
    )
    _write(
        out_dir / "group_load.csv",
        ["group_id", "group_name", "week", "slots"],
        [
            [gid, inst.groups.get(gid).name if gid in inst.groups else "", w, slots]
            for (gid, w), slots in sorted(group_load.items())
        ],
    )
    _write(
        out_dir / "room_util.csv",
        ["room_id", "room_name", "week", "slots"],
        [
            [rid, inst.rooms.get(rid).name if rid in inst.rooms else "", w, slots]
            for (rid, w), slots in sorted(room_load.items())
        ],
    )

# -------- ICS (one file per entity) --------

def _monday_of_week0(anchor: Optional[date] = None) -> date:
    if anchor is None:
        today = date.today()
        return today + timedelta(days=(7 - today.weekday()) % 7)
    return anchor - timedelta(days=anchor.weekday())

def _slot_dt(anchor_monday: date, day_name: str, slot_index: int, slot_minutes: int, slot_break_minutes: int, day_start: time, weeks_offset: int) -> tuple[datetime, datetime]:
    day_map = {name.upper()[:3]: i for i, name in enumerate(["MON","TUE","WED","THU","FRI","SAT","SUN"])}
    weekday = day_map.get(day_name.upper()[:3], 0)
    d = anchor_monday + timedelta(days=weekday + 7 * weeks_offset)
    start = datetime.combine(d, day_start) + timedelta(minutes=slot_index * (slot_minutes + slot_break_minutes))
    end = start + timedelta(minutes=slot_minutes)
    return start, end

def _ical_dt(dt: datetime) -> str:
    if dt.tzinfo is None:
        return dt.strftime("%Y%m%dT%H%M%S")
    return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")

def _ics_header(name: str) -> str:
    return "\n".join([
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Scheduling//Exporter//EN",
        f"X-WR-CALNAME:{name}",
    ])

def _ics_footer() -> str:
    return "END:VCALENDAR\n"

def _ics_event(uid: str, summary: str, start: datetime, end: datetime, location: str | None = None, description: str | None = None) -> str:
    lines = [
        "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{_ical_dt(datetime.now(timezone.utc))}",
        f"DTSTART:{_ical_dt(start)}",
        f"DTEND:{_ical_dt(end)}",
        f"SUMMARY:{summary}",
    ]
    if location:
        lines.append(f"LOCATION:{location}")
    if description:
        lines.append(f"DESCRIPTION:{description}")
    lines.append("END:VEVENT")
    return "\n".join(lines)

def _slug(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    s = re.sub(r"-+", "-", s).strip("-")
    return s or "x"

def _get_slot_minutes(inst: Instance) -> int:
    v = getattr(inst, "slot_minutes", None)
    return int(v) if isinstance(v, int) and v > 0 else _DEFAULT_SLOT_MINUTES

def _get_slot_break_minutes(inst: Instance) -> int:
    v = getattr(inst, "slot_break_minutes", None)
    return int(v) if isinstance(v, int) and v >= 0 else _DEFAULT_BREAK_MINUTES

def _get_day_start(inst: Instance) -> time:
    v = getattr(inst, "day_start_time", None)
    if isinstance(v, str) and ":" in v:
        try:
            h, m = map(int, v.split(":"))
            return time(h, m)
        except Exception:
            return _DEFAULT_START
    if isinstance(v, time):
        return v
    return _DEFAULT_START

def _write_ics(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")

def export_groups_ics_per_id(inst: Instance, schedule: Dict[int, Dict[str, Any]], out_dir: str,
                             week0_monday: Optional[date] = None) -> None:
    base = Path(out_dir)
    slot_minutes = _get_slot_minutes(inst)
    slot_break_minutes = _get_slot_break_minutes(inst)
    day_start = _get_day_start(inst)
    monday0 = _monday_of_week0(week0_monday)

    for gid, g in inst.groups.items():
        parts: List[str] = [_ics_header(f"Group {g.name}")]
        for a_id, info in schedule.items():
            if gid not in info.get("group_ids", []):
                continue
            w = int(info["week"])
            dname = str(info["day"])
            s0 = int(info["slot"])
            dur = int(info["duration"])
            # Emit one VEVENT per slot block of duration
            start, _ = _slot_dt(monday0, dname, s0, slot_minutes, slot_break_minutes, day_start, weeks_offset=w-1)
            end = start + timedelta(minutes=dur * slot_minutes)
            course = inst.courses.get(info["course_id"])
            room = inst.rooms.get(info.get("room_id")) if info.get("room_id") is not None else None
            staff = inst.staff.get(info.get("staff_id")) if info.get("staff_id") is not None else None
            summary = f"{course.code if course else ''} {info.get('kind','')}".strip()
            desc = f"Course: {course.name if course else ''} | Staff: {staff.name if staff else ''} | Groups: {','.join(str(x) for x in info.get('group_ids',[]))}"
            uid = f"group-{gid}-a{a_id}-w{w}-d{dname}-s{s0}"
            parts.append(_ics_event(uid, summary, start, end, location=(room.name if room else None), description=desc))
        parts.append(_ics_footer())
        name = f"group_{gid}_{_slug(g.name)}.ics"
        _write_ics("\n".join(parts), base / name)

def export_staff_ics_per_id(inst: Instance, schedule: Dict[int, Dict[str, Any]], out_dir: str,
                             week0_monday: Optional[date] = None) -> None:
    base = Path(out_dir)
    slot_minutes = _get_slot_minutes(inst)
    slot_break_minutes = _get_slot_break_minutes(inst)
    day_start = _get_day_start(inst)
    monday0 = _monday_of_week0(week0_monday)

    for sid, s in inst.staff.items():
        parts: List[str] = [_ics_header(f"Staff {s.name}")]
        for a_id, info in schedule.items():
            if sid != info.get("staff_id"):
                continue
            w = int(info["week"])
            dname = str(info["day"])
            s0 = int(info["slot"])
            dur = int(info["duration"])
            start, _ = _slot_dt(monday0, dname, s0, slot_minutes, slot_break_minutes, day_start, weeks_offset=w-1)
            end = start + timedelta(minutes=dur * slot_minutes)
            course = inst.courses.get(info["course_id"])
            room = inst.rooms.get(info.get("room_id")) if info.get("room_id") is not None else None
            summary = f"{course.code if course else ''} {info.get('kind','')}".strip()
            desc = f"Course: {course.name if course else ''} | Staff: {s.name} | Groups: {','.join(str(x) for x in info.get('group_ids',[]))}"
            uid = f"staff-{sid}-a{a_id}-w{w}-d{dname}-s{s0}"
            parts.append(_ics_event(uid, summary, start, end, location=(room.name if room else None), description=desc))
        parts.append(_ics_footer())
        name = f"staff_{sid}_{_slug(s.name)}.ics"
        _write_ics("\n".join(parts), base / name)

def export_rooms_ics_per_id(inst: Instance, schedule: Dict[int, Dict[str, Any]], out_dir: str,
                             week0_monday: Optional[date] = None) -> None:
    base = Path(out_dir)
    slot_minutes = _get_slot_minutes(inst)
    slot_break_minutes = _get_slot_break_minutes(inst)
    day_start = _get_day_start(inst)
    monday0 = _monday_of_week0(week0_monday)

    for rid, r in inst.rooms.items():
        parts: List[str] = [_ics_header(f"Room {r.name}")]
        for a_id, info in schedule.items():
            if rid != info.get("room_id"):
                continue
            w = int(info["week"])
            dname = str(info["day"])
            s0 = int(info["slot"])
            dur = int(info["duration"])
            start, _ = _slot_dt(monday0, dname, s0, slot_minutes, slot_break_minutes, day_start, weeks_offset=w-1)
            end = start + timedelta(minutes=dur * slot_minutes)
            course = inst.courses.get(info["course_id"])
            staff = inst.staff.get(info.get("staff_id")) if info.get("staff_id") is not None else None
            summary = f"{course.code if course else ''} {info.get('kind','')}".strip()
            desc = f"Course: {course.name if course else ''} | Staff: {staff.name if staff else ''} | Groups: {','.join(str(x) for x in info.get('group_ids',[]))}"
            uid = f"room-{rid}-a{a_id}-w{w}-d{dname}-s{s0}"
            parts.append(_ics_event(uid, summary, start, end, location=r.name, description=desc))
        parts.append(_ics_footer())
        name = f"room_{rid}_{_slug(r.name)}.ics"
        _write_ics("\n".join(parts), base / name)
